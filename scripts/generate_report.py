import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import tempfile

# ── Colors ──
BG = '#0a0a0a'
CARD = '#1a1a1a'
ACCENT = '#f43f5e'
ACCENT2 = '#fb923c'
TEXT = '#e5e5e5'
MUTED = '#a3a3a3'
BLUE = '#3b82f6'
GREEN = '#22c55e'
PURPLE = '#a855f7'
TEAL = '#14b8a6'

tmpdir = tempfile.mkdtemp()

def rounded_box(ax, x, y, w, h, text, color=ACCENT, fontsize=9, textcolor='white'):
    box = mpatches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                                   facecolor=color, edgecolor='white', linewidth=1.2, alpha=0.95)
    ax.add_patch(box)
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, color=textcolor, fontweight='bold', wrap=True)

def arrow(ax, x1, y1, x2, y2, color='white'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

def label_arrow(ax, x1, y1, x2, y2, label='', color='white'):
    arrow(ax, x1, y1, x2, y2, color)
    mx, my = (x1+x2)/2, (y1+y2)/2
    if label:
        ax.text(mx, my + 0.15, label, ha='center', va='bottom', fontsize=7, color=MUTED, style='italic')

# ═══════════════════════════════════════
# DIAGRAM 1: System Architecture
# ═══════════════════════════════════════
def create_system_architecture():
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('System Architecture', fontsize=16, color='white', fontweight='bold', pad=15)

    # Top - Guest
    rounded_box(ax, 3.8, 6.0, 2.4, 0.7, 'Hotel Guest\n(Web / Phone)', '#374151', 9)

    # Middle - Next.js App
    rounded_box(ax, 3.5, 4.5, 3.0, 0.7, 'Next.js 16 App\n(Frontend UI)', BLUE, 9)

    arrow(ax, 5.0, 6.0, 5.0, 5.25)

    # API Layer
    rounded_box(ax, 0.3, 2.8, 2.2, 0.7, '/api/stt\nWhisper STT', ACCENT, 9)
    rounded_box(ax, 3.9, 2.8, 2.2, 0.7, '/api/chat\nGPT-4o', ACCENT, 9)
    rounded_box(ax, 7.5, 2.8, 2.2, 0.7, '/api/telephony\nWebhook', ACCENT, 9)

    arrow(ax, 3.8, 4.5, 1.4, 3.55)
    arrow(ax, 5.0, 4.5, 5.0, 3.55)
    arrow(ax, 6.2, 4.5, 8.6, 3.55)

    # External Services
    rounded_box(ax, 0.3, 1.2, 2.2, 0.7, 'OpenAI\nWhisper API', PURPLE, 9)
    rounded_box(ax, 3.9, 1.2, 2.2, 0.7, 'OpenAI\nGPT-4o API', PURPLE, 9)
    rounded_box(ax, 7.5, 1.2, 2.2, 0.7, 'Response\nEngine', PURPLE, 9)

    arrow(ax, 1.4, 2.8, 1.4, 1.95)
    arrow(ax, 5.0, 2.8, 5.0, 1.95)
    arrow(ax, 8.6, 2.8, 8.6, 1.95)

    # Bottom - Config & Auth
    rounded_box(ax, 1.5, 0.1, 2.5, 0.7, 'Hotel Config\n(In-Memory Store)', TEAL, 8)
    rounded_box(ax, 6.0, 0.1, 2.5, 0.7, 'Auth System\nJWT + SQLite', GREEN, 8)

    arrow(ax, 5.0, 1.2, 2.75, 0.85)
    arrow(ax, 8.6, 1.2, 7.25, 0.85)

    path = os.path.join(tmpdir, 'architecture.png')
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor=BG)
    plt.close()
    return path

# ═══════════════════════════════════════
# DIAGRAM 2: User Voice Interaction Flow
# ═══════════════════════════════════════
def create_voice_flow():
    fig, ax = plt.subplots(1, 1, figsize=(10, 3.5))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis('off')
    ax.set_title('Voice Interaction Flow', fontsize=16, color='white', fontweight='bold', pad=15)

    # Flow boxes
    rounded_box(ax, 0.1, 1.2, 1.6, 0.9, 'Guest\nSpeaks', '#374151', 9)
    rounded_box(ax, 2.1, 1.2, 1.6, 0.9, 'Audio\nRecorded', BLUE, 9)
    rounded_box(ax, 4.1, 1.2, 1.6, 0.9, 'Whisper\nSTT', ACCENT, 9)
    rounded_box(ax, 6.1, 1.2, 1.6, 0.9, 'GPT-4o\nResponse', PURPLE, 9)
    rounded_box(ax, 8.1, 1.2, 1.6, 0.9, 'Speech\nSynthesis', GREEN, 9)

    label_arrow(ax, 1.75, 1.65, 2.1, 1.65, 'audio', 'white')
    label_arrow(ax, 3.75, 1.65, 4.1, 1.65, 'blob', 'white')
    label_arrow(ax, 5.75, 1.65, 6.1, 1.65, 'text', 'white')
    label_arrow(ax, 7.75, 1.65, 8.1, 1.65, 'reply', 'white')

    # Return arrow
    ax.annotate('', xy=(0.9, 1.15), xytext=(8.9, 0.6),
                arrowprops=dict(arrowstyle='->', color=ACCENT2, lw=1.8,
                                connectionstyle='arc3,rad=0.25'))
    ax.text(5.0, 0.25, 'Spoken response in guest\'s language', ha='center',
            fontsize=8, color=ACCENT2, style='italic', fontweight='bold')

    path = os.path.join(tmpdir, 'voice_flow.png')
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor=BG)
    plt.close()
    return path

# ═══════════════════════════════════════
# DIAGRAM 3: Admin Auth Flow
# ═══════════════════════════════════════
def create_auth_flow():
    fig, ax = plt.subplots(1, 1, figsize=(10, 3.5))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3.5)
    ax.axis('off')
    ax.set_title('Admin Authentication Flow', fontsize=16, color='white', fontweight='bold', pad=15)

    rounded_box(ax, 0.1, 1.2, 1.8, 0.9, 'Hotel Admin\nVisits /admin', '#374151', 9)
    rounded_box(ax, 2.5, 1.2, 1.8, 0.9, 'Register /\nLogin', BLUE, 9)
    rounded_box(ax, 4.9, 1.2, 1.8, 0.9, 'JWT Token\nin Cookie', ACCENT, 9)
    rounded_box(ax, 7.3, 1.2, 1.8, 0.9, 'proxy.ts\nGuard', PURPLE, 9)

    label_arrow(ax, 1.95, 1.65, 2.5, 1.65, 'credentials', 'white')
    label_arrow(ax, 4.35, 1.65, 4.9, 1.65, 'JWT issued', 'white')
    label_arrow(ax, 6.75, 1.65, 7.3, 1.65, 'validated', 'white')

    # Success arrow to settings
    rounded_box(ax, 7.3, 2.6, 1.8, 0.6, '/settings\nAccessible', GREEN, 9)
    arrow(ax, 8.2, 2.15, 8.2, 2.6)

    # Reject arrow
    ax.text(8.2, 0.7, 'Unauthorized?\nRedirect to /admin/login', ha='center',
            fontsize=7, color=ACCENT, style='italic')

    path = os.path.join(tmpdir, 'auth_flow.png')
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor=BG)
    plt.close()
    return path

# ═══════════════════════════════════════
# DIAGRAM 4: Tech Stack
# ═══════════════════════════════════════
def create_tech_stack():
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('Technology Stack', fontsize=16, color='white', fontweight='bold', pad=15)

    layers = [
        ('Frontend', ['Next.js 16', 'React 19', 'Tailwind v4', 'TypeScript'], BLUE),
        ('AI / APIs', ['GPT-4o', 'Whisper STT', 'Web Speech', 'TTS'], PURPLE),
        ('Backend', ['Route Handlers', 'proxy.ts', 'JWT Auth', 'bcryptjs'], ACCENT),
        ('Storage', ['SQLite', 'In-Memory', 'HTTP Cookies', 'better-sqlite3'], TEAL),
    ]

    for i, (label, items, color) in enumerate(layers):
        y = 3.8 - i * 1.1
        # Layer label
        ax.text(0.3, y + 0.25, label, fontsize=10, color=color, fontweight='bold', va='center')
        # Items
        for j, item in enumerate(items):
            x = 2.0 + j * 2.05
            rounded_box(ax, x, y, 1.8, 0.5, item, color, 8)

    path = os.path.join(tmpdir, 'tech_stack.png')
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor=BG)
    plt.close()
    return path

# ═══════════════════════════════════════
# BUILD THE DOCX
# ═══════════════════════════════════════
print("Generating diagrams...")
arch_path = create_system_architecture()
voice_path = create_voice_flow()
auth_path = create_auth_flow()
tech_path = create_tech_stack()
print("Diagrams generated.")

print("Building DOCX...")
doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

style_h1 = doc.styles['Heading 1']
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

style_h2 = doc.styles['Heading 2']
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0x24, 0x4, 0x4e)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()

def add_image(path, width=Inches(5.5)):
    doc.add_picture(path, width=width)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════
# CONTENT
# ════════════════════════════════════════════════

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\nInnovation Project Report\n\n')
run.font.size = Pt(28)
run.bold = True
run = p.add_run('StayNep AI Voice Receptionist\n')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xf4, 0x3f, 0x5e)
run = p.add_run('An Intelligent Multilingual Voice Assistant for Hotel Guest Services\n\n\n')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6b, 0x6b, 0x6b)

doc.add_page_break()

# ── Project Title ──
add_heading('Project Title')
add_para('StayNep AI Voice Receptionist — An Intelligent Multilingual Voice Assistant for Hotel Guest Services', bold=True)
doc.add_paragraph()

# ── Group Members ──
add_heading('Group Members')
add_table(
    ['S.N.', 'Name', 'Student ID'],
    [['1.', '[Member 1 Name]', '[Student ID]'],
     ['2.', '[Member 2 Name]', '[Student ID]'],
     ['3.', '[Member 3 Name]', '[Student ID]'],
     ['4.', '[Member 4 Name]', '[Student ID]']]
)

# ── Problem Statement ──
add_heading('Problem Statement')
add_para('The hospitality industry in Nepal and across South Asia faces a critical operational challenge: providing 24/7 multilingual guest support without incurring the high costs of round-the-clock multilingual staffing. Hotels, especially small to mid-sized properties, struggle with:')
add_bullet('Language barriers: Nepal\'s tourism industry serves guests from around the world, but most hotel staff are limited to Nepali and basic English, leading to miscommunication and poor guest experiences.')
add_bullet('24/7 availability: Front desk operations are often understaffed during night shifts and peak hours, causing long wait times and missed guest inquiries.')
add_bullet('Inconsistent service quality: Human receptionists vary in knowledge, mood, and availability. Guests frequently receive inaccurate or incomplete information about hotel policies, amenities, and local services.')
add_bullet('High operational costs: Hiring multilingual staff for 24/7 coverage is financially prohibitive for most hotels in Nepal, where average room rates are significantly lower than international standards.')
add_bullet('Scalability: As hotel chains grow, maintaining consistent service quality across multiple properties becomes increasingly difficult.')
add_para('These problems directly impact guest satisfaction scores, online reviews, and ultimately revenue — the lifeblood of the hospitality business.')
doc.add_paragraph()

# ── Proposed Solution ──
add_heading('Proposed Solution')
add_para('StayNep AI Voice Receptionist is a web-based, AI-powered voice assistant that acts as a virtual front desk receptionist for hotels. It provides instant, accurate, and multilingual responses to guest inquiries through natural voice conversation.')
doc.add_paragraph()
add_heading('How It Works', level=2)
add_bullet('Voice Input: Guests tap a microphone button and speak their question in any language. The audio is captured and sent to OpenAI Whisper, a state-of-the-art speech-to-text engine supporting 99+ languages.')
add_bullet('AI Processing: The transcribed text is processed by GPT-4o, configured with the specific hotel\'s data — room types, pricing, policies, amenities, dining options, and a custom AI persona.')
add_bullet('Voice Output: The response is spoken back to the guest using the browser\'s Speech Synthesis API in their preferred language, creating a natural conversational experience.')
add_bullet('Admin Dashboard: Hotel administrators can register, log in, and fully customize their hotel\'s configuration through a secure, authenticated admin panel.')
add_bullet('Telephony Integration: The system supports phone-based interactions via webhook integration with telephony providers (TingTing/Twilio).')
doc.add_paragraph()

# ── Innovative Idea ──
add_heading('Innovative Idea')
add_para('What makes StayNep AI Voice Receptionist unique:')
add_bullet('Voice-first, multilingual by design: Built from the ground up for voice interaction in 99+ languages. A Japanese tourist can speak in Japanese and receive a spoken response in Japanese — no typing, no translation apps.')
add_bullet('Zero-code hotel customization: Hotel owners with no technical background can fully configure the AI\'s knowledge base, personality, and branding through an intuitive admin dashboard.')
add_bullet('Dual-channel support (Web + Phone): The same AI brain serves both web-based voice interactions and traditional phone calls through telephony integration.')
add_bullet('Glassmorphic premium UI: Modern glassmorphism design with animated interactions, suitable for in-room tablet or lobby kiosk deployment.')
add_bullet('Nepal-focused with global capability: Built specifically for the Nepali hospitality market but architecturally designed to serve hotels worldwide.')
add_bullet('Integration-ready for hotel management systems: Designed as a modular component for the broader StayNep hotel management platform.')
doc.add_paragraph()

# ── SDG Goal Alignment ──
add_heading('SDG Goal Alignment')
add_heading('SDG 8: Decent Work and Economic Growth', level=2)
add_bullet('Enhances productivity of hotel operations by automating repetitive front-desk inquiries')
add_bullet('Enables small hotels in Nepal to compete with larger chains by providing world-class guest service at a fraction of the cost')
add_bullet('Supports tourism growth — a key economic driver for Nepal (contributing ~7.9% of GDP)')
add_heading('SDG 9: Industry, Innovation and Infrastructure', level=2)
add_bullet('Applies cutting-edge AI technology (GPT-4o, Whisper) to solve real hospitality industry challenges')
add_bullet('Builds digital infrastructure for Nepal\'s hotel industry, which is still largely manual')
add_bullet('Demonstrates how AI can be practically deployed in developing economies')
add_heading('SDG 11: Sustainable Cities and Communities', level=2)
add_bullet('Improves the tourist experience in Nepali cities, contributing to sustainable tourism')
add_bullet('Reduces the need for physical infrastructure through digital solutions')
doc.add_paragraph()

# ── MVP Features ──
add_heading('Minimum Viable Product (MVP) Features')
add_para('The current working prototype includes:')
add_bullet('Voice-to-voice conversation: Tap to speak, get spoken responses in real-time')
add_bullet('Multilingual support: Speech recognition and responses in 99+ languages via OpenAI Whisper + GPT-4o')
add_bullet('Hotel admin authentication: Secure registration and login system for hotel administrators')
add_bullet('Admin configuration dashboard with 8 customizable sections: Branding, Contact, Policies, Rooms, Dining, Amenities, Custom FAQ, AI Persona')
add_bullet('Route protection: Secure middleware (proxy.ts) preventing unauthorized access to admin features')
add_bullet('Telephony webhook: Phone call integration endpoint for voice-based phone interactions')
add_bullet('Concierge call mode: Direct telephony interface within the web app')
add_bullet('Responsive glassmorphic UI: Premium dark-themed interface with animations')
add_bullet('Conversation history: In-session chat log showing user and assistant messages')
doc.add_paragraph()

# ── Future Enhancements ──
add_heading('Future Additions/Enhancements')
add_table(
    ['#', 'Enhancement', 'Description'],
    [['1', 'Multi-tenant database', 'Persist each hotel\'s configuration in the database for true multi-hotel support'],
     ['2', 'Booking integration', 'Allow guests to check room availability and make reservations through voice'],
     ['3', 'StayNep integration', 'Connect voice assistant with the full StayNep hotel management system'],
     ['4', 'Analytics dashboard', 'Track guest inquiry patterns, peak hours, common questions'],
     ['5', 'Custom voice cloning', 'Allow hotels to create a unique AI voice matching their brand'],
     ['6', 'Offline mode', 'Cache common responses for areas with poor internet connectivity'],
     ['7', 'WhatsApp/Viber', 'Extend the AI to messaging platforms popular in Nepal'],
     ['8', 'Multi-property', 'Single admin account managing multiple hotel properties'],
     ['9', 'Guest feedback', 'Post-interaction satisfaction surveys'],
     ['10', 'Smart escalation', 'Automatically route complex queries to human staff']]
)

# ── Target Audience ──
add_heading('Target Audience')
add_table(
    ['Segment', 'Description'],
    [['Small to mid-sized hotels', 'Properties with 10-100 rooms that cannot afford 24/7 multilingual staff'],
     ['Tourist-area hotels', 'Hotels in Kathmandu, Pokhara, Chitwan, and Lumbini serving international tourists'],
     ['Hotel chains', 'Growing chains like StayNep needing consistent service across properties'],
     ['Boutique/heritage hotels', 'Properties seeking a premium tech-forward guest experience'],
     ['International tourists', 'Guests who need information in their native language'],
     ['Domestic travelers', 'Nepali-speaking guests who prefer voice interaction over text']]
)

# ── Technology/Tools Used ──
add_heading('Technology/Tools Used')
add_table(
    ['Category', 'Technology', 'Purpose'],
    [['Frontend', 'Next.js 16, React 19, TypeScript', 'Web application framework'],
     ['Styling', 'Tailwind CSS v4', 'Utility-first CSS with glassmorphism design'],
     ['Speech-to-Text', 'OpenAI Whisper API', 'Multilingual audio transcription (99+ languages)'],
     ['AI Chat', 'OpenAI GPT-4o', 'Intelligent, context-aware receptionist responses'],
     ['Text-to-Speech', 'Web Speech Synthesis API', 'Browser-native spoken responses'],
     ['Authentication', 'jose (JWT), bcryptjs', 'Secure token-based auth with password hashing'],
     ['Database', 'SQLite (better-sqlite3)', 'Lightweight hotel account storage'],
     ['Telephony', 'TingTing/Twilio webhooks', 'Phone call integration'],
     ['Version Control', 'Git, GitHub', 'Source code management'],
     ['Dev Tools', 'Turbopack, ESLint', 'Fast builds and code quality']]
)

# ── Tech Stack Diagram ──
add_image(tech_path, Inches(5.5))

# ── Timeline ──
add_heading('Project Timeline/Milestones')
add_table(
    ['Phase', 'Timeline', 'Milestone', 'Status'],
    [['Phase 1', 'Week 1-2', 'Project ideation, research, tech stack selection', 'Completed'],
     ['Phase 2', 'Week 3-4', 'Core voice assistant with AI integration', 'Completed'],
     ['Phase 3', 'Week 5-6', 'Premium UI, multilingual support, telephony', 'Completed'],
     ['Phase 4', 'Week 7', 'Authentication system, admin security', 'Completed'],
     ['Phase 5', 'Week 8', 'Migration to OpenAI (Whisper + GPT-4o)', 'Completed'],
     ['Phase 6', 'Week 9-10', 'Testing, bug fixes, documentation', 'In Progress'],
     ['Phase 7', 'Week 11-12', 'Presentation preparation and demo', 'Upcoming'],
     ['Phase 8', 'Post-submission', 'StayNep platform integration', 'Planned']]
)

# ── Potential Challenges ──
add_heading('Potential Challenges')
add_table(
    ['Challenge', 'Impact', 'Mitigation Strategy'],
    [['API costs', 'Per-request charges scale with users', 'Response caching, cheaper models for simple queries, rate limits'],
     ['Internet dependency', 'Nepal has inconsistent internet', 'Offline fallback with cached FAQs, browser Speech API as fallback'],
     ['Audio quality', 'Background noise affects accuracy', 'Noise suppression, text input fallback, Whisper language hints'],
     ['Latency', 'Voice needs sub-2s response time', 'Streaming responses, optimized prompts, edge deployment'],
     ['Data privacy', 'Guest voice data sensitivity', 'No audio storage, HTTPS only, data retention policies'],
     ['Language accuracy', 'Low-resource languages less accurate', 'Prioritize testing top tourist languages (Nepali, Hindi, Chinese, Japanese)'],
     ['Adoption resistance', 'Staff fear of job displacement', 'Position as augmentation tool for overflow and after-hours']]
)

# ── Market Research ──
add_heading('Market Research')
add_heading('Nepal Hotel Industry Overview', level=2)
add_bullet('Nepal has 1,200+ registered hotels and 3,500+ accommodation providers')
add_bullet('Tourism contributed NPR 240 billion (~$1.8B USD) to Nepal\'s economy in 2024')
add_bullet('Nepal welcomed 1.2 million+ international tourists in 2025, with a government target of 2.5 million by 2030')
add_bullet('68% of tourists cite "communication difficulty" as a challenge during their Nepal stay')
doc.add_paragraph()

add_heading('Competitive Landscape', level=2)
add_table(
    ['Solution', 'Limitation'],
    [['Traditional reception', 'Limited hours, language barriers, inconsistent quality'],
     ['Generic chatbots (Tidio, Drift)', 'Text-only, English-centric, no hotel-specific knowledge'],
     ['IVR phone systems', 'Rigid menu trees, frustrating UX, no AI'],
     ['Connie by Hilton', 'Enterprise-only, prohibitively expensive for Nepali hotels']]
)

add_heading('Revenue Model (Post-College)', level=2)
add_bullet('Freemium: Basic voice assistant free for up to 100 interactions/month')
add_bullet('Pro: NPR 2,000/month (~$15) for unlimited interactions + analytics')
add_bullet('Enterprise: Custom pricing for hotel chains with multi-property support')
doc.add_paragraph()

# ── Visual Representations ──
add_heading('Visual Representations')
doc.add_paragraph()

add_heading('System Architecture Diagram', level=2)
add_image(arch_path, Inches(5.8))
doc.add_paragraph()

add_heading('Voice Interaction Flow', level=2)
add_image(voice_path, Inches(5.8))
doc.add_paragraph()

add_heading('Admin Authentication Flow', level=2)
add_image(auth_path, Inches(5.8))

# ── Footer ──
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nThis project is developed as part of our college innovation initiative and serves as the foundation for StayNep — a comprehensive hotel management platform for Nepal\'s growing hospitality industry.')
run.font.size = Pt(10)
run.italic = True
run.font.color.rgb = RGBColor(0x6b, 0x6b, 0x6b)

# ── Save ──
output = '/Users/prabinsharma/Desktop/StayNep_Innovation_Project_Report.docx'
doc.save(output)
print(f"Report saved to: {output}")

# Cleanup temp images
for f in [arch_path, voice_path, auth_path, tech_path]:
    os.remove(f)
os.rmdir(tmpdir)
